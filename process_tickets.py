import pandas as pd 
from docx import Document
from google import genai 
from dotenv import load_dotenv
load_dotenv()
client= genai.Client()
mt = pd.read_csv('data/tickets.csv')## MT-main table
total_tickets=len(mt) 
unique_days = mt["Submission_Date"].nunique()
avg_tickets_per_day=round(total_tickets /unique_days, 1 )
priority_list =  (mt["Priority_Level"].value_counts(normalize=True) * 100).round(1).to_dict()
top_category = mt["Issue_Category"].value_counts().idxmax()
top_category_count = mt["Issue_Category"].value_counts().max()
avg_satisfaction =round(mt["Satisfaction_Score"].mean(), 2)
top_3_agents = mt.groupby("Assigned_Agent")["Satisfaction_Score"].mean().nlargest(3).round(2).to_dict()
channels_list = (mt["Ticket_Channel"].value_counts(normalize=True)*100).round(1).to_dict()
promt = f"""
You are a Senior Customer Support Analyst. Analyze the following aggregated support ticket data and generate a concise business report:

- Avrage tickets per day: {avg_tickets_per_day}
- Priority distribution: {priority_list}
- Most critical category: {top_category}
- Average customer satisfaction score (1-5): {avg_satisfaction}
- Top 3 performing agents: {top_3_agents}
- Support ticket channels: {channels_list}
Provide 3 key takeaways and 2 actionable recommendations to improve customer support operations.
"""
response= client.models.generate_content(
    model="gemini-3.1-flash-lite", contents=promt
)
doc=Document()
doc.add_heading("Customer Support Analytics & AI Report", level=1)
doc.add_heading("1. Aggregated Key Metrics", level=2)
doc.add_paragraph(f"Average tickets per day: {avg_tickets_per_day}")
doc.add_paragraph(f"Priority distribution (%): {priority_list}")
doc.add_paragraph(
    f"Top issue category: {top_category} ({top_category_count} tickets)"
)
doc.add_paragraph(f"Average Satisfaction Score: {avg_satisfaction} / 5")
doc.add_paragraph(f"Top 3 agents by score: {top_3_agents}")
doc.add_paragraph(f"Channel distribution (%): {channels_list}")
doc.add_paragraph(response.text)
doc.save("support_report.docx")
print("Report successfully generated and saved to support_report.docx!")